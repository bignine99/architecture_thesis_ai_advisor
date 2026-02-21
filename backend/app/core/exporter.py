"""
분석 결과를 TXT, PDF, DOCX 형식으로 내보내는 모듈
"""
import os
import io
import re
from datetime import datetime


def export_to_txt(review: str, guidelines: str, filename: str) -> bytes:
    """분석 결과를 TXT 파일로 변환한다."""
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    
    content = f"""{'='*70}
  Architectural Thesis Advisor AI - 분석 보고서
  파일: {filename}
  생성일시: {now}
{'='*70}

{'─'*70}
  논문 평가 종합평
{'─'*70}

{review}

{'─'*70}
  실행 과제 및 개선 전략
{'─'*70}

{guidelines}

{'='*70}
  본 보고서는 AI 기반 건축학 논문 비평 시스템에 의해 자동 생성되었습니다.
{'='*70}
"""
    return content.encode('utf-8-sig')


def export_to_pdf(review: str, guidelines: str, filename: str) -> bytes:
    """분석 결과를 PDF 파일로 변환한다."""
    from fpdf import FPDF
    
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    
    # Load Korean font
    font_path = r"C:\Windows\Fonts\malgun.ttf"
    bold_font_path = r"C:\Windows\Fonts\malgunbd.ttf"
    
    font_name = "Helvetica"  # fallback
    
    if os.path.exists(font_path):
        try:
            pdf.add_font("Malgun", "", font_path)
            font_name = "Malgun"
            if os.path.exists(bold_font_path):
                pdf.add_font("MalgunB", "", bold_font_path)
        except Exception as e:
            print(f"Font loading warning: {e}")
    
    # --- Title Page ---
    pdf.add_page()
    pdf.ln(50)
    
    if font_name == "Malgun":
        pdf.set_font("MalgunB" if os.path.exists(bold_font_path) else "Malgun", "", 22)
    else:
        pdf.set_font(font_name, "B", 22)
    
    pdf.set_text_color(50, 50, 80)
    pdf.cell(0, 15, "논문 평가 종합평", align="C", new_x="LMARGIN", new_y="NEXT")
    
    pdf.ln(5)
    pdf.set_font(font_name, "", 12)
    pdf.set_text_color(100, 100, 120)
    pdf.cell(0, 10, "Critical Peer Review Report", align="C", new_x="LMARGIN", new_y="NEXT")
    
    pdf.ln(20)
    pdf.set_font(font_name, "", 10)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 8, f"분석 대상: {filename}", align="C", new_x="LMARGIN", new_y="NEXT")
    now_str = datetime.now().strftime("%Y년 %m월 %d일 %H:%M")
    pdf.cell(0, 8, f"생성일시: {now_str}", align="C", new_x="LMARGIN", new_y="NEXT")
    
    # --- Review ---
    pdf.add_page()
    _write_section_to_pdf(pdf, font_name, "논문 평가 종합평", review)
    
    # --- Guidelines ---
    pdf.add_page()
    _write_section_to_pdf(pdf, font_name, "실행 과제 및 개선 전략", guidelines)
    
    return bytes(pdf.output())


def _write_section_to_pdf(pdf, font_name: str, title: str, text: str):
    """섹션 제목과 본문을 PDF에 작성한다."""
    page_w = 190  # usable page width in mm (A4 210 - 10 left - 10 right)
    
    # Title
    pdf.set_font(font_name, "", 14)
    pdf.set_text_color(40, 40, 80)
    pdf.multi_cell(page_w, 10, title)
    pdf.set_draw_color(150, 150, 200)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(6)
    
    # Clean markdown bold markers for PDF
    clean_text = text.replace('**', '')
    
    # Write text line by line
    lines = clean_text.split('\n')
    for line in lines:
        stripped = line.strip()
        
        if not stripped:
            pdf.ln(3)
            continue
        
        # Headings
        if stripped.startswith('### '):
            pdf.ln(3)
            pdf.set_font(font_name, "", 11)
            pdf.set_text_color(70, 70, 140)
            pdf.multi_cell(page_w, 6, stripped[4:])
            pdf.ln(2)
            continue
        elif stripped.startswith('## '):
            pdf.ln(4)
            pdf.set_font(font_name, "", 12)
            pdf.set_text_color(50, 50, 100)
            pdf.multi_cell(page_w, 7, stripped[3:])
            pdf.ln(2)
            continue
        elif stripped.startswith('# '):
            pdf.ln(5)
            pdf.set_font(font_name, "", 14)
            pdf.set_text_color(40, 40, 80)
            pdf.multi_cell(page_w, 8, stripped[2:])
            pdf.ln(3)
            continue
        
        # Horizontal rule
        if stripped in ('---', '***', '___'):
            pdf.set_draw_color(200, 200, 200)
            y = pdf.get_y()
            pdf.line(10, y + 2, 200, y + 2)
            pdf.ln(5)
            continue
        
        # Table rows - render as plain text
        if '|' in stripped and stripped.startswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if all(re.match(r'^[-:]+$', c) for c in cells):
                continue
            row_text = "  ".join(cells)
            pdf.set_font(font_name, "", 9)
            pdf.set_text_color(60, 60, 60)
            pdf.multi_cell(page_w, 5, row_text)
            continue
        
        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            pdf.set_font(font_name, "", 10)
            pdf.set_text_color(50, 50, 50)
            pdf.multi_cell(page_w, 6, "  " + stripped[2:])
            continue
        
        # Numbered list
        if re.match(r'^\d+\.\s', stripped):
            pdf.set_font(font_name, "", 10)
            pdf.set_text_color(50, 50, 50)
            pdf.multi_cell(page_w, 6, "  " + stripped)
            continue
        
        # Regular text
        pdf.set_font(font_name, "", 10)
        pdf.set_text_color(50, 50, 50)
        pdf.multi_cell(page_w, 6, stripped)


def export_to_docx(review: str, guidelines: str, filename: str) -> bytes:
    """분석 결과를 DOCX 파일로 변환한다."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # --- Title ---
    title = doc.add_heading('논문 평가 종합평', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Critical Peer Review Report')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 120)
    
    # Meta info
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = meta.add_run(f'\n분석 대상: {filename}\n')
    run1.font.size = Pt(10)
    now_str = datetime.now().strftime("%Y년 %m월 %d일 %H:%M")
    run2 = meta.add_run(f'생성일시: {now_str}')
    run2.font.size = Pt(10)
    
    doc.add_page_break()
    
    # --- Review Section ---
    doc.add_heading('논문 평가 종합평', level=1)
    _write_markdown_to_docx(doc, review)
    
    doc.add_page_break()
    
    # --- Guidelines Section ---
    doc.add_heading('실행 과제 및 개선 전략', level=1)
    _write_markdown_to_docx(doc, guidelines)
    
    # --- Footer ---
    doc.add_paragraph('')
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('본 보고서는 AI 기반 건축학 논문 비평 시스템에 의해 자동 생성되었습니다.')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _write_markdown_to_docx(doc, text: str):
    """마크다운 텍스트를 DOCX 문서에 렌더링한다."""
    from docx.shared import Pt
    
    lines = text.split('\n')
    table_rows = []
    in_table = False
    
    for line in lines:
        stripped = line.strip()
        
        # Table handling
        if '|' in stripped and stripped.startswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if all(re.match(r'^[-:]+$', c) for c in cells if c):
                continue  # Skip separator row
            table_rows.append(cells)
            in_table = True
            continue
        elif in_table and table_rows:
            _flush_table(doc, table_rows)
            table_rows = []
            in_table = False
        
        if not stripped:
            continue
        
        # Headings
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
            continue
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
            continue
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
            continue
        
        # Horizontal rule
        if stripped in ('---', '***', '___'):
            continue
        
        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            content = stripped[2:].replace('**', '')
            doc.add_paragraph(content, style='List Bullet')
            continue
        
        # Numbered lists
        if re.match(r'^\d+\.\s', stripped):
            content = re.sub(r'^\d+\.\s', '', stripped).replace('**', '')
            doc.add_paragraph(content, style='List Number')
            continue
        
        # Regular paragraph
        content = stripped.replace('**', '')
        p = doc.add_paragraph(content)
        for run in p.runs:
            run.font.size = Pt(10)
    
    # Flush remaining table
    if table_rows:
        _flush_table(doc, table_rows)


def _flush_table(doc, rows: list):
    """표 데이터를 DOCX 테이블로 추가한다."""
    from docx.shared import Pt
    
    if not rows:
        return
    
    max_cols = max(len(r) for r in rows)
    
    try:
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        
        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                if j < max_cols:
                    cell = table.rows[i].cells[j]
                    cell.text = cell_text.replace('**', '')
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9)
                            if i == 0:
                                run.bold = True
    except Exception as e:
        # Fallback: add as plain text
        for row_data in rows:
            doc.add_paragraph(" | ".join(r.replace('**', '') for r in row_data))
    
    doc.add_paragraph('')  # Space after table
